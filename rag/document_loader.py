"""Load text from PDF, DOCX and TXT files with basic validation."""
import os
from pypdf import PdfReader
import docx

MAX_FILE_BYTES = 20 * 1024 * 1024  # 20 MB per file


def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def _read_docx(path: str) -> str:
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:          # capture tabular text too
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


LOADERS = {".pdf": _read_pdf, ".docx": _read_docx, ".txt": _read_txt}


def load_document(path: str) -> str:
    """Return the extracted text of a supported document, or raise ValueError."""
    ext = os.path.splitext(path)[1].lower()
    if ext not in LOADERS:
        raise ValueError(f"Unsupported file type '{ext}'. Allowed: PDF, DOCX, TXT.")
    size = os.path.getsize(path)
    if size > MAX_FILE_BYTES:
        raise ValueError(
            f"'{os.path.basename(path)}' is too large "
            f"({size / 1e6:.1f} MB). Max is {MAX_FILE_BYTES / 1e6:.0f} MB."
        )
    text = LOADERS[ext](path)
    if not text or not text.strip():
        raise ValueError(
            f"No extractable text found in '{os.path.basename(path)}' "
            "(it may be empty or a scanned image)."
        )
    return text
