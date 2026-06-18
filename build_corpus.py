"""Build the India Constitution corpus (PDF / DOCX / TXT) from sourced text.

Sources:
  * Preamble: Yash-Handa/The_Constitution_Of_India (COI.json)
  * Part III / IV / IVA articles: `IndianConstitution` PyPI package data
Bare constitutional text is freely reproducible (Copyright Act, 1957, s. 52).
Article 45 is overridden with its current (86th Amendment, 2002) text.
"""
import json
import os

OUT = os.path.join(os.path.dirname(__file__), "data")
os.makedirs(OUT, exist_ok=True)

# --- load sources ---
# Requires: pip install IndianConstitution reportlab python-docx
import indianconstitution
PKG = os.path.join(os.path.dirname(indianconstitution.__file__), "data", "constitution.json")
by = {str(x["article"]): x for x in json.load(open(PKG))}

PREAMBLE = (
    "WE, THE PEOPLE OF INDIA, having solemnly resolved to constitute India into a "
    "SOVEREIGN SOCIALIST SECULAR DEMOCRATIC REPUBLIC and to secure to all its citizens:\n"
    "JUSTICE, social, economic and political;\n"
    "LIBERTY of thought, expression, belief, faith and worship;\n"
    "EQUALITY of status and of opportunity;\n"
    "and to promote among them all\n"
    "FRATERNITY assuring the dignity of the individual and the unity and integrity of the Nation;\n"
    "IN OUR CONSTITUENT ASSEMBLY this twenty-sixth day of November, 1949, do HEREBY ADOPT, "
    "ENACT AND GIVE TO OURSELVES THIS CONSTITUTION."
)

# Article 45 current text (substituted by the 86th Amendment, 2002)
by["45"] = {
    "article": "45",
    "title": "Provision for early childhood care and education to children below the age of six years",
    "description": ("The State shall endeavour to provide early childhood care and education for "
                    "all children until they complete the age of six years."),
}

DISCLAIMER = ("Note: Text compiled from public sources for demonstration of a document-chat "
              "system. For the authoritative, current text refer to indiacode.nic.in.\n")

# --- Part III: Fundamental Rights (curated, current articles) ---
PART_III = [
    (None, ["12", "13"]),
    ("Right to Equality", ["14", "15", "16", "17", "18"]),
    ("Right to Freedom", ["19", "20", "21", "21A", "22"]),
    ("Right against Exploitation", ["23", "24"]),
    ("Right to Freedom of Religion", ["25", "26", "27", "28"]),
    ("Cultural and Educational Rights", ["29", "30"]),
    ("Right to Constitutional Remedies", ["32"]),
]

# --- Part IV: Directive Principles ---
PART_IV = ["36", "37", "38", "39", "39A", "40", "41", "42", "43", "43A",
           "44", "45", "46", "47", "48", "48A", "49", "50", "51"]


def art_line(num):
    x = by[num]
    return f"Article {num}. {x['title']}.\n{x['description'].strip()}"


# ---------- 1) Fundamental Rights -> PDF ----------
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

styles = getSampleStyleSheet()
h1 = ParagraphStyle("h1", parent=styles["Title"], fontSize=18, spaceAfter=10)
h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, spaceBefore=10, spaceAfter=4)
art = ParagraphStyle("art", parent=styles["BodyText"], fontSize=10.5, leading=15, spaceAfter=8)
small = ParagraphStyle("small", parent=styles["BodyText"], fontSize=8, textColor="#666666")


def esc(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")


pdf_path = os.path.join(OUT, "Fundamental-Rights.pdf")
doc = SimpleDocTemplate(pdf_path, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm,
                        topMargin=2 * cm, bottomMargin=2 * cm)
flow = [Paragraph("Constitution of India — Part III: Fundamental Rights", h1),
        Paragraph(esc(DISCLAIMER), small), Spacer(1, 6)]
for heading, nums in PART_III:
    if heading:
        flow.append(Paragraph(heading, h2))
    for n in nums:
        flow.append(Paragraph(esc(art_line(n)), art))
doc.build(flow)

# ---------- 2) Directive Principles -> DOCX ----------
import docx

d = docx.Document()
d.add_heading("Constitution of India — Part IV: Directive Principles of State Policy", level=1)
d.add_paragraph(DISCLAIMER.strip()).italic = True
for n in PART_IV:
    x = by[n]
    d.add_heading(f"Article {n}. {x['title']}", level=2)
    d.add_paragraph(x["description"].strip())
docx_path = os.path.join(OUT, "Directive-Principles.docx")
d.save(docx_path)

# ---------- 3) Preamble + Fundamental Duties -> TXT ----------
txt_path = os.path.join(OUT, "Preamble-and-Fundamental-Duties.txt")
with open(txt_path, "w", encoding="utf-8") as f:
    f.write("CONSTITUTION OF INDIA\n")
    f.write(DISCLAIMER + "\n")
    f.write("THE PREAMBLE\n\n")
    f.write(PREAMBLE + "\n\n")
    f.write("=" * 60 + "\n\n")
    f.write("PART IVA — FUNDAMENTAL DUTIES\n\n")
    f.write(art_line("51A") + "\n")

for p in (pdf_path, docx_path, txt_path):
    print(f"wrote {p} ({os.path.getsize(p)} bytes)")
